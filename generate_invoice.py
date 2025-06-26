import os
import subprocess
from pathlib import Path
import datetime
from docx import Document

def convert_docx_to_pdf(docx_path,output_dir):
    command = [
        r'C:\Program Files\LibreOffice\program\soffice.exe',
        '--headless',
        '--convert-to', 'pdf',
        '--outdir', str(output_dir),
        str(docx_path)
    ]
    subprocess.run(command,check=True)

def batch_convert_docx2pdf(folder_with_docx,output_folder):
    folder = Path(folder_with_docx)
    output = Path(output_folder)
    output.mkdir(exist_ok=True)

    for docx_file in folder.glob('*.docx'):
        convert_docx_to_pdf(docx_file,output)
        os.remove(docx_file)


def fill_the_data(name,address,number,email,due_date,service,amt,discount,invoice_no):
    global inv_count
    tax = ((amt-discount)*18)/100
    doc = Document('Templates/basic-invoice.docx')

    # Data to fill
    invoice_data = {
        "table-1":
        {
            "[Company Name]": "Shiv Shakti Pvt. Ltd.",
            "invoice_num": f"{invoice_no}",                     # Replace Invoice #
            "Current Date": f"{datetime.date.today()}",           # Replace date
            "Company_address": "121, Cannon Street 1",                 # Replace address
            "Zip Code": "Surat, Gujarat, 395007",
            "phone_no": "+91-9912348765", # Replace number
            "Due_Date": f"{due_date}"
        },
        'table-2':
        {
            "[Name]": f"{name}",
            "[Street Address]": f"{address}",
            "Bill_Phone": f"{number}",
            "Bill_Email": f"{email}",
        },
        'table-3':
        {
            "Service": f"{service}",
            "Amt": f"{amt:.2f}",
            "Dis": f"{discount:.2f}",
            "Taxation": f"{tax:.2f}",
            "Totl": f"{((amt-discount) + tax):.2f}"
        }
    }

    help_inquiry = {
        "Name": "P. Shreyash",
        "Phone": '+91-9977654231',
        "email@address.com": "shivshakti.info@example.com"
    }

    for para in doc.paragraphs:
        for key, value in help_inquiry.items():
            if key in para.text:
                # We need to handle runs separately to preserve formatting
                for run in para.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, value)

    def replace_placeholders_in_table(table, replacements):
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        for key, value in replacements.items():
                            if key in run.text:
                                run.text = run.text.replace(key, value)

    # Apply replacements for each table
    for i, key in enumerate(["table-1", "table-2", "table-3"]):
        replace_placeholders_in_table(doc.tables[i], invoice_data[key])

    # Save the filled invoice
    doc.save(f'Invoices/Invoice_{invoice_no}.docx')

    print(f'{email} --> Bill Generated Successfully.')



